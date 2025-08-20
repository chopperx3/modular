from pathlib import Path
from typing import Dict, Tuple
import re
from docx import Document
from docx.shared import Pt

OUT_DIR = Path(__file__).resolve().parent / "out" / "renewed"
OUT_DIR.mkdir(parents=True, exist_ok=True)

GEN_DATE = r"(?:[0-3]?\d[\/\.-][01]?\d[\/\.-](?:\d{2}|\d{4}))"
GEN_MONEY = r"(?:\$?\s?(?:\d{1,3}(?:[.,]\d{3})*|\d+)(?:[.,]\d{2})?)"
GEN_EMAIL = r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}"
GEN_PHONE = r"(?:\+?\d{1,3}[\s-]?)?(?:\(?\d{2,3}\)?[\s-]?)?\d{3,4}[\s-]?\d{3,4}"
GEN_URL = r"https?://[^\s]+"

def normalize_text(t: str) -> str:
    t = (t or "").replace("\r", "")
    t = t.replace("“", '"').replace("”", '"').replace("’", "'")
    t = re.sub(r"(\w)-\n(\w)", r"\1\2", t)
    t = re.sub(r"\s+([,.;:!?])", r"\1", t)
    paras = [re.sub(r"\s*\n\s*", " ", p).strip() for p in re.split(r"\n{2,}", t)]
    t = "\n\n".join([p for p in paras if p])
    return re.sub(r"[ \t]{2,}", " ", t).strip()

def parse_generic_fields(text: str) -> Dict[str, str]:
    u = text.upper(); out: Dict[str, str] = {}
    def first(p): 
        m = re.search(p, u, re.I); 
        return m.group(0) if m else ""
    if d:=first(GEN_DATE):  out["Fecha"]=d
    if m:=first(GEN_MONEY): out["Monto"]=m
    if e:=first(GEN_EMAIL): out["Email"]=e
    if p:=first(GEN_PHONE): out["Teléfono"]=p
    if u:=first(GEN_URL):   out["URL"]=u
    return out

def build_docx(clean_text: str, filename: str, doc_type_id: int | None,
               fields: Dict[str, str], result_id: int) -> Tuple[Path, Path]:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    doc.add_heading("Documento renovado", level=0)
    p = doc.add_paragraph(); r = p.add_run("Archivo: "); r.bold=True; p.add_run(filename)
    if doc_type_id is not None:
        p = doc.add_paragraph(); r = p.add_run("Tipo: "); r.bold=True; p.add_run(str(doc_type_id))
    if fields:
        doc.add_paragraph(); doc.add_heading("Campos", level=1)
        table = doc.add_table(rows=1, cols=2)
        hdr = table.rows[0].cells; hdr[0].text="Clave"; hdr[1].text="Valor"
        for k,v in fields.items():
            c=table.add_row().cells; c[0].text=str(k); c[1].text=str(v)
    doc.add_paragraph(); doc.add_heading("Contenido", level=1)
    for para in clean_text.split("\n\n"): doc.add_paragraph(para)
    stem = f"renewed_{result_id}"
    docx_path = OUT_DIR / f"{stem}.docx"
    txt_path  = OUT_DIR / f"{stem}.txt"
    doc.save(docx_path); txt_path.write_text(clean_text, encoding="utf-8")
    return docx_path, txt_path

def renew_document(*, text: str, filename: str, doc_type_id: int | None, result_id: int):
    clean = normalize_text(text)
    fields = parse_generic_fields(clean)
    docx_path, txt_path = build_docx(clean, filename, doc_type_id, fields, result_id)
    preview = clean[:400] + ("…" if len(clean) > 400 else "")
    return docx_path, txt_path, fields, preview
